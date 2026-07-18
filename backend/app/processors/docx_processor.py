"""
Enterprise RAG AI Assistant — DOCX Processor
=============================================
Extracts text and metadata from DOCX files using ``python-docx``.

Extraction strategy
-------------------
- Iterate over ``document.paragraphs`` to capture body text and headings.
- Iterate over ``document.tables`` → rows → cells for tabular content.
- Join paragraphs with newlines, tables with double-newlines.
- Extract core properties (title, author, subject, description).

Page count for DOCX is not natively available without a running Word instance.
We estimate it as ``max(1, total_words // 250)`` which is a commonly used
approximation (250 words per page).

Dependencies: python-docx, lxml
"""

from __future__ import annotations

from pathlib import Path

from app.processors import normalizer
from app.processors.base_processor import BaseProcessor, ProcessingResult
from app.processors.language_detector import detect_language


class DOCXProcessor(BaseProcessor):
    """Extracts text and metadata from DOCX files via python-docx."""

    _WORDS_PER_PAGE = 250  # Approximate page size for page count estimation

    def _extract(self, path: Path) -> ProcessingResult:
        """
        Open the DOCX, extract paragraphs + tables, normalise, detect language.

        Args:
            path: Path to the DOCX file.

        Returns:
            ``ProcessingResult`` with all fields populated.

        Raises:
            RuntimeError: If the file cannot be read as a valid DOCX.
        """
        from docx import Document  # type: ignore[import-untyped]
        from docx.opc.exceptions import (
            PackageNotFoundError,  # type: ignore[import-untyped]
        )

        try:
            doc = Document(str(path))
        except (PackageNotFoundError, Exception) as exc:
            raise RuntimeError(f"Failed to open DOCX at {path}: {exc}") from exc

        # ── 1. Body paragraphs (includes headings) ────────────────────────────
        paragraph_lines: list[str] = []
        for para in doc.paragraphs:
            text = para.text
            if text and text.strip():
                paragraph_lines.append(text)

        body_text = "\n".join(paragraph_lines)

        # ── 2. Table content ──────────────────────────────────────────────────
        table_parts: list[str] = []
        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                table_parts.append("\n".join(rows))

        table_text = "\n\n".join(table_parts)

        raw_text = "\n\n".join(part for part in [body_text, table_text] if part.strip())

        # ── 3. Core properties metadata ───────────────────────────────────────
        meta: dict[str, str] = {}
        try:
            props = doc.core_properties
            for attr in ("title", "author", "subject", "description", "keywords"):
                value = getattr(props, attr, None)
                if value and isinstance(value, str) and value.strip():
                    meta[attr] = value.strip()
        except Exception:
            pass  # metadata is best-effort

        # ── 4. Normalise and compute statistics ───────────────────────────────
        clean_text = normalizer.normalize(raw_text)
        wc = normalizer.word_count(clean_text)
        page_count = max(1, wc // self._WORDS_PER_PAGE) if wc > 0 else 1
        language = detect_language(clean_text)

        return ProcessingResult(
            raw_text=raw_text,
            clean_text=clean_text,
            language=language,
            page_count=page_count,
            word_count=wc,
            character_count=normalizer.character_count(clean_text),
            processing_time=0.0,
            metadata=meta,
        )
